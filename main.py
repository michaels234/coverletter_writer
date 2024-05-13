import os
from dotenv import load_dotenv
import requests
#import json
from docx import Document
from bs4 import BeautifulSoup
from openai import OpenAI
from google.colab import drive
from weasyprint import HTML
import pyperclip
drive.mount('/content/drive')

copied_job = """copied job description here
"""


job_url = 'https://www.dice.com/job-detail/example'
job_url = 'cybercoders'  # for overwriting to just use copied_job
if job_url not in ['cybercoders']:
  domain = job_url.split("://")[1].split(".")
  if domain[0] == "www":
      domain = domain[1]
  else:
      domain = domain[0]
else:
    domain = job_url

cover_letter_instructions = """Please write me a cover letter based on the information provided.
Key points to remember when writing the letter:
1. make it 250 words or less.
2. don't use any markdown formatting, just plain text.
3. don't include addresses at the beginning, just start with Dear Hiring Team.
4. don't mention any cities or countries by name in the letter, especially don't mention Tokyo or Japan.
5. in the intro paragraph be sure to mention my 7 years of experience in software engineering.
6. be sure to mention how many years of experience i have in the relevant places for anything that I have 3 or more years of experience in. If I have less than 3 years of experience in something, just refer to it qualitatively.
7. don't bring up a piece of experience or one of my projects that isn't relevant to the job, and vice-versa don't mention parts of the job that I don't seem to have experience in.
8. don't use strong exaggerated language like 'I am deeply passionate about...'.
9. if there's a technology in the job description that I've only used in my personal projects as listed at the bottom of my resume, refer to it briefly, just basically stating that I've used it in personal projects.
10. if a job entails anything related to math or physics or machines, make sure to highlight my strength in those thanks to my background on mathematics, physics, and mechanical engineering. Notably, any quant jobs I want to be sure to mention my strong skill and background in mathematics and statistics."""

interest_instructions = """Please write me a short 150 word explanation for 'why i'm interested in working for this company'.
Key points to remember when writing the letter:
1. don't use any markdown formatting, just plain text.
2. don't mention any cities or countries by name in the letter, especially don't mention Tokyo or Japan.
3. don't use strong exaggerated language like 'I am deeply passionate about...'."""

amazing_fit_instructions = """Please write me an explanation for 'why I would be an amazing fit' for this job.
Key points to remember when writing the explanation:
1. Maximum 1000 characters.
2. don't use any markdown formatting, just plain text.
3. don't mention any cities or countries by name in the letter, especially don't mention Tokyo or Japan.
4. don't use strong exaggerated language like 'I am deeply passionate about...'."""

instructions = {"dice": cover_letter_instructions, "wellfound": interest_instructions, "cybercoders": amazing_fit_instructions}

accenture_projects = """1. I did full-stack development for multiple in-store web apps and mobile apps for a cosmetics company. The app is for users in the store to browse cosmetics products, as well as virtually put cosmetics on their face with a camera display that uses facial recognition to apply makeup on their face like a filter. I led small sub-teams on this project, owning features from start to finish. We used an external company for the AI makeup camera and implemented it into our app. The stack was python run in aws lambda, mysql, with api gateway and rds among other aws services, javascript, typescript, html, css frontend with reactive javascript frameworks, particularly vue.js. mobile was android, java, kotlin. This project is the one I mentioned in my resume "Increased avg. user spending by $200/yr".
2. I did development on Android TV system settings apps for an electronics company. The stack was android, java, kotlin, sqlite. i had to lead overseas teams on this project in 2 languages (japanese and English). This project is the one I mentioned in my resume as "Completed 6mo of work in 4 on distressed project.".
3. I developed an android app from scratch, using a generative AI api we paid for that takes an image and text as input and uses image recognition and generative ai to understand the image and the input text which is usually a question about the image, and then it uses generative ai to output an answer, thats all the api we paid for, but i implemented that api into an android app for use by blind people so they can get information about their surroundings. they can ask a question and get information about what the camera sees. it was an android wearable device with a camera speaker and mic. it spoke the answers aloud. Android, kotlin ,java. This project is the one I mentioned in my resume as "Inspired 6-figure project contract".
4. I worked full-stack on an iot system that has a front and backend, its for a tire manufacturer for use in the factory, the system uses receivers around the factory and transmitters on mobile equipment, and using those iot signals the app estimates equipment location so factory workers don't have to look for equipment as much. This one increased factory efficiency for this process. Python AWS rds dynamodb lambda event bridge batch iot JavaScript vue.js reactive framework. We used AI for taking signals and outputting estimated location. We trained with tensorflow. This project is the one I mentioned in my resume as "Improved efficiency of one factory process by 75%".
5. I created a whole app myself for use on a farm that employs disabled employees for disabled community outreach. The app helps managers on the farm check seeds planted by using image recognition to see if the seeds are planted too deep or too shallow. It decreased time managers had to check the workers tasks. Python flask JavaScript vue.js opencv and we used sagemaker to train the ai. This project is the one I mentioned in my resume as "Reduced time managers spent checking workers’ tasks by 65%"."""

class CoverLetterWriter():
    def __init__(self):
        try:
            if domain in ["cybercoders"]:
                self.job = copied_job
            else:
                self.job = self.scrape_job()
        except CaughtException as e:
            print(e)
            return

        try:
            self.resume = self.read_resume()
        except CaughtException as e:
            print(e)
            return

        try:
            self.cover_letter = self.first_draft()
        except CaughtException as e:
            print(e)

        try:
            self.write_letter()
        except CaughtException as e:
            print(e)

        return

    def scrape_job(self):
        response = requests.get(job_url)

        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            job = ''

            if domain == "dice":
                job_title_element = soup.find('h1', attrs={'data-cy': 'jobTitle'})
                if job_title_element:
                    job_title = job_title_element.text.strip()
                else:
                    raise CaughtException("job_title doesn't exist")
                job_description_element = soup.find('section', class_='job-description')
                if job_description_element:
                    job_description = job_description_element.get_text(separator='\n', strip=True)
                else:
                    raise CaughtException("job_description doesn't exist")

                job = job_title + '\n' + job_description
            elif domain == "wellfound":
                target_element = soup.find('div', {'data-test': 'JobListing'})
                if target_element:
                    job = target_element.get_text()
                else:
                    raise CaughtException("target_element doesn't exist")
            else:
                raise CaughtException("can't get job description from this domain")
            return job
        else:
            raise CaughtException("scrape_job response failure.", response)
        return

    def read_resume(self):
        file_path = "/content/drive/MyDrive/Resume_Michael_Simeone.docx"
        try:
            res = Document(file_path)
            resume = ''
            for paragraph in res.paragraphs:
                resume += paragraph.text + '\n'
        except Exception as e:
            raise CaughtException("read_resume failure.", e)
        return resume

    def first_draft(self):
        text = f"""
Below are My Resume, My List of Accenture Projects, and a Job Description I'm interested in applying for.
{instructions[domain]}


My Resume:
{self.resume}


My List of Accenture Projects:
{accenture_projects}


Job Description:
{self.job}
        """

        try:
            load_dotenv("/content/drive/MyDrive/openai.env")
            client = OpenAI()
            response = client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {
                        "role": "user",
                        "content": text
                    }
                ],
                temperature=1,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )
            plain_text = response.choices[0].message.content
            if domain == 'dice':
                html = f"""<html><body><div style="font-family: Helvetica, Arial, sans-serif; font-size: 12px; white-space: pre-wrap">{plain_text}</div></body></html>"""
            elif domain in ['wellfound', 'cybercoders']:
                html = plain_text
            else:
                raise CaughtException("domain doesn't match anything prepared.")
            return html
        except Exception as e:
            raise CaughtException("failure getting response in first_draft.", e)

        return

    def write_letter(self):
        try:
            # letter = Document()
            # letter.add_paragraph(self.cover_letter)
            # letter.save('Cover_Letter.docx')
            if domain == 'dice':
                pdf_path = "/content/drive/MyDrive/cover_letter.pdf"
                HTML(string=self.cover_letter).write_pdf(pdf_path)
            elif domain in ['wellfound', 'cybercoders']:
                print(self.cover_letter)
                # pyperclip.copy(self.cover_letter)
            else:
                raise CaughtException("domain doesn't match anything prepared.")
        except Exception as e:
            raise CaughtException("write_letter failure.", e)
        return


class CaughtException(Exception):
    def __init__(self, message):
        super().__init__(f"CustomException: {message}")


CoverLetterWriter()